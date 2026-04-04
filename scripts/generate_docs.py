import os
from docx import Document
from fpdf import FPDF

def generate_word(md_file, docx_file):
    doc = Document()
    doc.add_heading('AfriBiz - Documentation', 0)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].replace('**', '')
            doc.add_paragraph(text, style='List Bullet')
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            text = line.replace('**', '')
            doc.add_paragraph(text)
            
    doc.save(docx_file)
    print(f"Successfully generated {docx_file}")

def generate_pdf(md_file, pdf_file):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Use standard Arial (no special characters)
    pdf.set_font("Arial", 'B', 20)
    pdf.cell(200, 20, txt="AfriBiz - Documentation", ln=True, align='L')
    pdf.ln(5)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip().encode('latin-1', 'replace').decode('latin-1')
        if not line:
            pdf.ln(2)
            continue
            
        if line.startswith('# '):
            pdf.set_font("Arial", 'B', 18)
            pdf.ln(5)
            pdf.cell(200, 10, txt=line[2:], ln=True, align='L')
            pdf.ln(2)
        elif line.startswith('## '):
            pdf.set_font("Arial", 'B', 15)
            pdf.ln(3)
            pdf.cell(200, 10, txt=line[3:], ln=True, align='L')
            pdf.ln(1)
        elif line.startswith('### '):
            pdf.set_font("Arial", 'B', 13)
            pdf.ln(2)
            pdf.cell(200, 10, txt=line[4:], ln=True, align='L')
        elif line.startswith('- ') or line.startswith('* '):
            pdf.set_font("Arial", '', 11)
            text = "  - " + line[2:].replace('**', '')
            pdf.multi_cell(0, 8, txt=text)
        elif line.startswith('---'):
            pdf.add_page()
        else:
            pdf.set_font("Arial", '', 11)
            text = line.replace('**', '')
            pdf.multi_cell(0, 8, txt=text)
            
    pdf.output(pdf_file)
    print(f"Successfully generated {pdf_file}")

if __name__ == "__main__":
    md_path = r"e:\New folder\DOCUMENTATION.md"
    docx_path = r"e:\New folder\AfriBiz_Documentation.docx"
    pdf_path = r"e:\New folder\AfriBiz_Documentation.pdf"
    
    generate_word(md_path, docx_path)
    generate_pdf(md_path, pdf_path)
